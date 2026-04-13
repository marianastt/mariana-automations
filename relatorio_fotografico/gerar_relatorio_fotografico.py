#!/usr/bin/env python3
"""
Gera relatório fotográfico de auditoria de operador Polen.
Puxa dados do BigQuery e gera documento Word (.docx) seguindo o manual de marca Polen 2026.

Uso:
  python3 gerar_relatorio_fotografico.py --name jandaia
  python3 gerar_relatorio_fotografico.py --cnpj 21977543000133
  python3 gerar_relatorio_fotografico.py --subsidiary-id <UUID>

Opções:
  --fotos-dir DIR     Diretório com fotos baixadas
  --storage-url URL   URL base do storage para baixar fotos
  --output FILE       Nome do arquivo de saída
"""

import argparse
import json
import os
import re
import subprocess
import sys
import tempfile
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml
except ImportError:
    print("Instalando python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import OxmlElement, parse_xml

BQ_CLI = "/opt/homebrew/bin/bq"
BQ_PROJECT = "analytics-big-query-242119"
BQ_DATASET = "selo_polen_prod_gcp_fivetran_public"

# BigQuery client (preferred over bq CLI)
_bq_client = None
def _get_bq_client():
    global _bq_client
    if _bq_client is None:
        try:
            from google.cloud import bigquery as _bq
            _bq_client = _bq.Client(project=BQ_PROJECT)
        except Exception:
            _bq_client = False
    return _bq_client if _bq_client else None

# ── Paleta Polen (Manual de Marca 2026) ──
AZUL_ESCURO = RGBColor(0x25, 0x3D, 0x55)      # #253D55
GRAFITE     = RGBColor(0x60, 0x60, 0x60)       # #606060
CINZA_CLARO = "F2F2F2"                          # #F2F2F2 (hex string for shading)
BRANCO      = "FFFFFF"
AZUL_HEX    = "253D55"

FONT_NAME = "Red Hat Display"

# Caminhos do logo Polen
LOGO_WHITE_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "polen_logo_white.png")
LOGO_DARK_PATH = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "polen_logo_dark.png")

# Seções padrão de fotos
PHOTO_SECTIONS = [
    ("Fotos da fachada",
     "Fachada e entrada do empreendimento"),
    ("Fotos da área de triagem",
     "Área de triagem de materiais"),
    ("Fotos do galpão de triagem",
     "Galpão de triagem"),
    ("Fotos dos equipamentos da área de triagem (prensas, balanças, esteiras)",
     "Equipamentos: prensas, balanças, esteiras"),
    ("Fotos da prensa (se houver)",
     "Prensa"),
    ("Fotos da Balança (se houver)",
     "Balança"),
    ("Fotos da esteira/mesa triagem",
     "Esteira / Mesa de triagem"),
    ("Fotos da empilhadeira (se houver)",
     "Empilhadeira"),
    ("Fotos dos funcionários na operação/utilização de EPI",
     "Funcionários na operação / Utilização de EPI"),
    ("Foto dos veículos (se houver)",
     "Veículos"),
]


# ═══════════════════════════════════════════════════════════════
#  HELPERS - BigQuery
# ═══════════════════════════════════════════════════════════════

def bq_query(sql):
    # Try Python BigQuery client first (no gcloud dependency)
    client = _get_bq_client()
    if client:
        try:
            rows = list(client.query(sql).result())
            return [dict(r) for r in rows]
        except Exception as e:
            print(f"  [WARN] BQ client error: {e}")

    # Fallback to bq CLI
    cmd = [BQ_CLI, "query", "--format=json", "--max_rows=1000", "--nouse_legacy_sql", sql]
    result = subprocess.run(cmd, capture_output=True, text=True)
    output = result.stdout.strip()
    lines = [l for l in output.split("\n") if not l.startswith("Waiting on")]
    output = "\n".join(lines).strip()
    if not output or output == "[]":
        return []
    try:
        return json.loads(output)
    except json.JSONDecodeError:
        start = output.find("[")
        end = output.rfind("]")
        if start >= 0 and end > start:
            try:
                return json.loads(output[start:end+1])
            except json.JSONDecodeError:
                pass
        print(f"  [WARN] BQ parse error: {output[:200]}")
        return []


def find_operator(args):
    if args.subsidiary_id:
        where = f"s.id = '{args.subsidiary_id}'"
    elif args.cnpj:
        cnpj_clean = re.sub(r'\D', '', args.cnpj)
        where = f"REPLACE(REPLACE(REPLACE(s.cnpj, '.', ''), '/', ''), '-', '') = '{cnpj_clean}'"
    elif args.name:
        where = f"LOWER(c.name) LIKE '%{args.name.lower()}%'"
    else:
        print("Erro: forneça --subsidiary-id, --cnpj ou --name")
        sys.exit(1)

    rows = bq_query(f"""
        SELECT s.id as subsidiary_id, s.cnpj, s.type, s.state as uf, s.city,
               c.name as company_name, c.id as company_id
        FROM `{BQ_PROJECT}.{BQ_DATASET}.subsidiaries` s
        JOIN `{BQ_PROJECT}.{BQ_DATASET}.companies` c ON s.companyid = c.id
        WHERE c.stakeholder = 'operator' AND s._fivetran_deleted IS NOT TRUE AND {where}
    """)
    if not rows:
        print("Operador não encontrado.")
        sys.exit(1)
    if len(rows) > 1:
        print(f"Múltiplos operadores ({len(rows)}):")
        for r in rows:
            print(f"  - {r['company_name']} | {r['cnpj']} | {r['subsidiary_id']}")
        sys.exit(1)
    return rows[0]


def get_audit(subsidiary_id):
    rows = bq_query(f"""
        SELECT * FROM `{BQ_PROJECT}.{BQ_DATASET}.audits`
        WHERE subsidiaryid = '{subsidiary_id}'
        AND _fivetran_deleted IS NOT TRUE AND status = 'Approved'
        ORDER BY COALESCE(aprovedat, createdat) DESC LIMIT 1
    """)
    if not rows:
        print("Nenhuma auditoria aprovada encontrada.")
        sys.exit(1)
    return rows[0]


def get_observations(subsidiary_id):
    """Busca observações de TODAS as auditorias (pega o valor mais recente não-vazio)."""
    rows = bq_query(f"""
        SELECT materialinputcontrolnote, operationobservation,
               noteonequipment, notes
        FROM `{BQ_PROJECT}.{BQ_DATASET}.audits`
        WHERE subsidiaryid = '{subsidiary_id}'
        AND _fivetran_deleted IS NOT TRUE AND status = 'Approved'
        ORDER BY COALESCE(aprovedat, createdat) DESC
    """)
    obs_fields = {
        "materialinputcontrolnote": "Controle de Entrada de Materiais",
        "operationobservation": "Observação Operacional",
        "noteonequipment": "Observação sobre Equipamentos",
        "notes": "Observações Gerais",
    }
    result = []
    seen = set()
    for row in rows:
        for field, label in obs_fields.items():
            if field not in seen and row.get(field) and row[field].strip():
                result.append((label, row[field].strip()))
                seen.add(field)
    return result


def get_address(audit):
    dp = audit.get("dataprofile", "{}")
    if isinstance(dp, str):
        try:
            dp = json.loads(dp)
        except:
            dp = {}
    rua = dp.get("logradouro", "")
    numero = dp.get("numero", "")
    complemento = dp.get("complemento", "")
    bairro = dp.get("bairro", "")
    cidade = dp.get("municipio", "")
    uf = dp.get("uf", "")
    cep = dp.get("cep", "")
    parts = [rua]
    if numero:
        parts[0] += f", {numero}"
    if complemento:
        parts[0] += f" - {complemento}"
    if bairro:
        parts.append(bairro)
    cidade_uf = f"{cidade}/{uf}" if cidade and uf else cidade or uf
    if cidade_uf:
        parts.append(cidade_uf)
    if cep:
        parts.append(f"CEP {cep}")
    return " - ".join(parts) if parts[0] else "Endereço não informado"


def get_all_attachments(subsidiary_id):
    return bq_query(f"""
        SELECT aa.storagelocationkey, aa.filename, aa.contenttype,
          IFNULL(ac.name, '(sem categoria)') as category_name,
          IFNULL(ac.category, '(sem)') as category_group,
          a.id as audit_id
        FROM `{BQ_PROJECT}.{BQ_DATASET}.auditattachments` aa
        JOIN `{BQ_PROJECT}.{BQ_DATASET}.audits` a ON a.id = aa.auditid
        LEFT JOIN `{BQ_PROJECT}.{BQ_DATASET}.attachment_categories` ac ON ac.id = aa.attachment_category_id
        WHERE a.subsidiaryid = '{subsidiary_id}'
        AND a._fivetran_deleted IS NOT TRUE AND a.status = 'Approved'
        AND aa._fivetran_deleted IS NOT TRUE AND aa.contenttype LIKE 'image/%'
        ORDER BY a.aprovedat DESC, ac.category, aa.filename
    """)


def get_answers(audit_id):
    rows = bq_query(f"""
        SELECT q.name as question, q.category, a.value as answer
        FROM `{BQ_PROJECT}.{BQ_DATASET}.answers` a
        JOIN `{BQ_PROJECT}.{BQ_DATASET}.questions` q ON q.id = a.question_id
        WHERE a.audit_id = '{audit_id}'
        AND a._fivetran_deleted IS NOT TRUE AND q._fivetran_deleted IS NOT TRUE
        ORDER BY q.category, q.name
    """)
    # Carrega o dicionário de opções (UUID → label) para perguntas tipo Select
    opts = bq_query(f"""
        SELECT DISTINCT qd.id, qd.name
        FROM `{BQ_PROJECT}.{BQ_DATASET}.question_dependents` qd
        JOIN `{BQ_PROJECT}.{BQ_DATASET}.answers` a ON a.question_id = qd.question_id
        WHERE a.audit_id = '{audit_id}' AND qd._fivetran_deleted IS NOT TRUE
    """)
    global _options_dict
    _options_dict = {o['id']: o['name'] for o in opts}
    return rows


# Dicionário global de opções (UUID → label), populado por get_answers
_options_dict = {}


# ═══════════════════════════════════════════════════════════════
#  HELPERS - Formatação Word
# ═══════════════════════════════════════════════════════════════

def set_font(run, size=11, bold=False, color=None, italic=False):
    """Aplica fonte Red Hat Display a um run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Força Red Hat Display no XML (necessário para alguns leitores)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)


def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(doc, text, size=11, bold=False, color=None, italic=False,
                         align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=Pt(6)):
    """Adiciona parágrafo com fonte Polen, justificado por padrão."""
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_section_title(doc, text, level=1):
    """Título de seção com barra lateral azul."""
    if level == 1:
        # Barra azul + título grande
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        # Coluna da barra azul (estreita)
        bar_cell = table.cell(0, 0)
        bar_cell.width = Cm(0.3)
        set_cell_shading(bar_cell, AZUL_HEX)
        bar_cell.text = ""

        # Coluna do texto
        text_cell = table.cell(0, 1)
        text_cell.width = Cm(15.7)
        text_cell.text = ""
        p = text_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(f"  {text}")
        set_font(run, size=18, bold=True, color=AZUL_ESCURO)

        # Remover bordas da tabela
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{ border_name}')
            border.set(qn('w:val'), 'none')
            border.set(qn('w:sz'), '0')
            borders.append(border)
        tblPr.append(borders)

        doc.add_paragraph()  # espaço
        return table

    else:
        # Subtítulo com linha inferior
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_font(run, size=14, bold=True, color=AZUL_ESCURO)

        # Linha separadora fina
        sep = doc.add_paragraph()
        sep.paragraph_format.space_before = Pt(0)
        sep.paragraph_format.space_after = Pt(8)
        pPr = sep._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:color'), AZUL_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

        return p


def build_info_table(doc, label_values):
    """Tabela de informações estilizada com cores Polen."""
    table = doc.add_table(rows=len(label_values), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Definir largura das colunas
    for row_idx, (label, value) in enumerate(label_values):
        cell_l = table.cell(row_idx, 0)
        cell_l.width = Cm(7)
        cell_l.text = ""
        run = cell_l.paragraphs[0].add_run(label)
        set_font(run, size=10, bold=True, color=AZUL_ESCURO)
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell_l, CINZA_CLARO)

        cell_v = table.cell(row_idx, 1)
        cell_v.width = Cm(9)
        cell_v.text = ""
        run = cell_v.paragraphs[0].add_run(str(value) if value else "—")
        set_font(run, size=10)
        cell_v.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Estilizar bordas da tabela
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'D0D0D0')
        borders.append(border)
    tblPr.append(borders)

    doc.add_paragraph()
    return table


def add_photo_placeholder(doc, legend, photo=None, fotos_dir=None, storage_url=None):
    """Box de foto com legenda."""
    foto_path = None

    if photo and fotos_dir:
        key = photo["storagelocationkey"]
        for candidate in [
            os.path.join(fotos_dir, key),
            os.path.join(fotos_dir, os.path.basename(key)),
            os.path.join(fotos_dir, photo["filename"]),
        ]:
            if os.path.exists(candidate):
                foto_path = candidate
                break

    if photo and not foto_path and storage_url:
        key = photo["storagelocationkey"]
        url = f"{storage_url.rstrip('/')}/{key}"
        try:
            import urllib.request
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(key)[1])
            urllib.request.urlretrieve(url, tmp.name)
            foto_path = tmp.name
        except Exception as e:
            print(f"  Erro ao baixar {photo['filename']}: {e}")

    if foto_path:
        try:
            doc.add_picture(foto_path, width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except:
            _add_placeholder_box(doc, legend)
    else:
        _add_placeholder_box(doc, legend)

    # Legenda centralizada
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(12)
    run = cap.add_run(legend)
    set_font(run, size=9, bold=True, color=GRAFITE, italic=True)


def _add_placeholder_box(doc, legend):
    """Box cinza claro como placeholder de foto."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CINZA_CLARO)

    # Altura do box (~7cm)
    tr = cell._tc.getparent()
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '3969')
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)

    # Largura
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), '8505')
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Bordas sutis
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for bname in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{bname}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'D0D0D0')
        borders.append(b)
    tblPr.append(borders)

    run_space = p.add_run("\n\n")
    set_font(run_space, size=11)
    run = p.add_run(f"[ Inserir foto: {legend} ]")
    set_font(run, size=11, color=GRAFITE, italic=True)


def yn(val):
    if val is None:
        return "—"
    if isinstance(val, str):
        return "Sim" if val.lower() == "true" else "Não"
    return "Sim" if val else "Não"


def clean_answer(val):
    if not val:
        return "—"
    val = val.strip('"')
    if val == "undefined":
        return "—"
    # UUID único — busca rótulo na tabela de opções (Sim/Não/Mensal/Anual/etc.)
    if re.match(r'^[0-9a-f]{8}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{4}-[0-9a-f]{12}$', val):
        return _options_dict.get(val, val)
    # Array de UUIDs — multi-select, junta os rótulos
    if val.startswith('['):
        try:
            arr = json.loads(val.replace('\\"', '"'))
            if all(re.match(r'^[0-9a-f]{8}-', str(x)) for x in arr):
                labels = [_options_dict.get(x, x) for x in arr]
                return ", ".join(labels) if labels else "—"
        except:
            pass
    return val


MESES_PT = {1:"janeiro", 2:"fevereiro", 3:"março", 4:"abril", 5:"maio", 6:"junho",
            7:"julho", 8:"agosto", 9:"setembro", 10:"outubro", 11:"novembro", 12:"dezembro"}


def format_date(date_str):
    if not date_str:
        return "—"
    if isinstance(date_str, datetime):
        return f"{date_str.day} de {MESES_PT[date_str.month]} de {date_str.year}"
    try:
        dt = datetime.strptime(str(date_str)[:10], "%Y-%m-%d")
        return f"{dt.day} de {MESES_PT[dt.month]} de {dt.year}"
    except:
        return str(date_str)[:10]


# ═══════════════════════════════════════════════════════════════
#  GERAÇÃO DO RELATÓRIO
# ═══════════════════════════════════════════════════════════════

def answers_dict(answers):
    """Converte lista de answers em dict {question: answer}."""
    d = {}
    for a in answers:
        d[a["question"]] = a["answer"].strip('"') if a["answer"] else ""
    return d


def generate_report(operator, audit, all_attachments, answers, observations, args):
    doc = Document()

    # ── Configuração da página ──
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Fonte padrão do documento
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(11)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)

    company_name = operator["company_name"]
    cnpj = operator["cnpj"]
    uf = (operator.get("uf") or "").strip() or "—"
    operator_type = "Cooperativa" if operator["type"] == "Coop" else "Empresa"
    address = get_address(audit)
    approval_date = format_date(audit.get("aprovedat", ""))
    mode = getattr(args, "mode", "both")  # "dco" | "fotografico" | "both"

    # ════════════════════════════════════════
    #  CAPA
    # ════════════════════════════════════════

    # Logo Polen (versão escura em fundo branco)
    logo_path = LOGO_DARK_PATH if os.path.exists(LOGO_DARK_PATH) else LOGO_WHITE_PATH
    if os.path.exists(logo_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(20)
        p_logo.paragraph_format.space_after = Pt(0)
        run_logo = p_logo.add_run()
        run_logo.add_picture(logo_path, width=Cm(7))

    # Linha separadora fina azul abaixo do logo
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(8)
    sep.paragraph_format.space_after = Pt(0)
    pPr = sep._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom_border = OxmlElement('w:bottom')
    bottom_border.set(qn('w:val'), 'single')
    bottom_border.set(qn('w:sz'), '6')
    bottom_border.set(qn('w:color'), AZUL_HEX)
    pBdr.append(bottom_border)
    pPr.append(pBdr)

    for _ in range(4):
        doc.add_paragraph()

    # Título capa
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RELATÓRIO FOTOGRÁFICO")
    set_font(run, size=28, bold=True, color=AZUL_ESCURO)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("E")
    set_font(run, size=14, color=GRAFITE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DECLARAÇÃO DE CAPACIDADE OPERACIONAL")
    set_font(run, size=22, bold=True, color=AZUL_ESCURO)

    for _ in range(4):
        doc.add_paragraph()

    # Box operador na capa — 3 pares (OPERADOR, CNPJ, ESTADO)
    op_table = doc.add_table(rows=6, cols=1)
    op_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cover_fields = [
        ("OPERADOR", company_name),
        ("CNPJ", cnpj),
        ("ESTADO", uf),
    ]
    for i, (label, value) in enumerate(cover_fields):
        # Label (linha azul)
        cell_label = op_table.cell(i * 2, 0)
        cell_label.text = ""
        set_cell_shading(cell_label, AZUL_HEX)
        p = cell_label.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_font(run, size=12, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        # Valor (linha cinza)
        cell_val = op_table.cell(i * 2 + 1, 0)
        cell_val.text = ""
        set_cell_shading(cell_val, CINZA_CLARO)
        p = cell_val.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(value)
        set_font(run, size=13, bold=True, color=AZUL_ESCURO)

    # Borda da tabela
    tbl = op_table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), AZUL_HEX)
        borders.append(b)
    tblPr.append(borders)

    for _ in range(3):
        doc.add_paragraph()

    # Barra inferior azul
    bar_table2 = doc.add_table(rows=1, cols=1)
    bar_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar_cell2 = bar_table2.cell(0, 0)
    set_cell_shading(bar_cell2, AZUL_HEX)
    bar_cell2.text = ""
    tr2 = bar_cell2._tc.getparent()
    trPr2 = tr2.get_or_add_trPr()
    trH2 = OxmlElement('w:trHeight')
    trH2.set(qn('w:val'), '567')
    trH2.set(qn('w:hRule'), 'exact')
    trPr2.append(trH2)
    tbl2 = bar_table2._tbl
    tblPr2 = tbl2.tblPr if tbl2.tblPr is not None else OxmlElement('w:tblPr')
    borders2 = OxmlElement('w:tblBorders')
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{bn}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        borders2.append(b)
    tblPr2.append(borders2)

    doc.add_page_break()

    # ════════════════════════════════════════
    #  CONTEÚDO DO QUESTIONÁRIO (DCO)
    # ════════════════════════════════════════
    ans = answers_dict(answers)

    if mode != "fotografico":
        if mode == "dco":
            intro_title = "DECLARAÇÃO DE CAPACIDADE OPERACIONAL"
            intro_text = (
                f"A seguir é apresentada a DECLARAÇÃO DE CAPACIDADE OPERACIONAL "
                f"do Operador de Reciclagem {company_name}, inscrita sob o CNPJ {cnpj}."
            )
        else:
            intro_title = "RELATÓRIO FOTOGRÁFICO"
            intro_text = (
                f"A seguir é apresentado o RELATÓRIO FOTOGRÁFICO do Operador de Reciclagem "
                f"{company_name}, inscrita sob o CNPJ {cnpj}."
            )
        add_section_title(doc, intro_title, level=1)
        add_styled_paragraph(doc, intro_text)
        doc.add_paragraph()

    # ── Informações Gerais ──
    if mode != "fotografico":
        add_section_title(doc, "Informações Gerais", level=2)
        build_info_table(doc, [
            ("Razão Social", company_name),
            ("CNPJ", cnpj),
            ("Tipo", operator_type),
            ("Endereço", address),
            ("Data de Aprovação da Auditoria", approval_date),
        ])

        # Placeholder para foto de localização (abaixo das informações gerais)
        p_sv = doc.add_paragraph()
        p_sv.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_sv = p_sv.add_run("Localização do Empreendimento")
        set_font(run_sv, size=13, bold=True, color=AZUL_ESCURO)
        _add_placeholder_box(doc, f"Foto Google Street View — {address}")
        cap_sv = doc.add_paragraph()
        cap_sv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_sv.paragraph_format.space_after = Pt(12)
        run_cap_sv = cap_sv.add_run(address)
        set_font(run_cap_sv, size=9, bold=True, color=GRAFITE, italic=True)
        doc.add_paragraph()

    # ── Administrativo ──
    admin_items = []
    emp_h = ans.get("Homens", "")
    emp_m = ans.get("Mulheres", "")
    if emp_h or emp_m:
        total = int(emp_h or 0) + int(emp_m or 0)
        admin_items.append(("Funcionários", f"{total} (Homens: {emp_h} / Mulheres: {emp_m})"))
    div_admin = ans.get("Possui divisão administrativa definida? (Presidente, conselheiro, tesoureiro)", "")
    if div_admin and not re.match(r'^[0-9a-f]{8}-', div_admin):
        admin_items.append(("Divisão Administrativa", div_admin))
    renda = ans.get("Renda Média Por Pessoa", "")
    if renda:
        admin_items.append(("Renda Média por Pessoa (R$)", renda))
    if admin_items and mode != "fotografico":
        add_section_title(doc, "Administrativo", level=2)
        build_info_table(doc, admin_items)

    # ── Equipamentos ──
    # Mostra TODOS os equipamentos, mesmo com quantidade 0
    if mode != "fotografico":
        equip_items = []
        for label, key in [("Prensas", "Prensas"), ("Balanças", "Balanças"),
                            ("Mesas de Triagem", "Mesas"), ("Esteiras", "Esteiras"),
                            ("Empilhadeira", "Empilhadeira"), ("Veículos", "Veículos")]:
            val = ans.get(key, "")
            equip_items.append((label, val if val else "0"))
        add_section_title(doc, "Equipamentos", level=2)
        build_info_table(doc, equip_items)

    # ── Infraestrutura ──
    infra_items = []
    infra_questions = [
        ("Local coberto?", "Local Coberto"),
        ("Possuem Refeitório?", "Refeitório"),
        ("Área de triagem possui ventiladores/ventilação adequada?", "Ventilação Adequada"),
    ]
    for q_key, label in infra_questions:
        val = clean_answer(ans.get(q_key, ""))
        if val and val != "—":
            infra_items.append((label, val))
    if infra_items and mode != "fotografico":
        add_section_title(doc, "Infraestrutura", level=2)
        build_info_table(doc, infra_items)

    # ── Segurança ──
    seg_items = []
    seg_questions = [
        ("Os funcionários usam EPIs?", "Uso de EPI"),
        ("Há treinamento para a utilização dos EPIs?", "Treinamento EPI"),
        ("Frequência em que é realizado o treinamento para a utilização dos EPIs", "Frequência Treinamento EPI"),
        ("Iluminação adequada?", "Iluminação Adequada"),
        ("Extintores dentro da validade?", "Extintores Válidos"),
        ("Saídas de emergência desobstruídas?", "Saídas de Emergência"),
        ("Há controle sobre os níveis de ruídos gerados?", "Controle de Ruído"),
        ("Há somente Funcionários / Associados / Cooperados maiores de 18 anos?", "Somente Maiores de 18 anos"),
    ]
    for q_key, label in seg_questions:
        val = clean_answer(ans.get(q_key, ""))
        if val and val != "—":
            seg_items.append((label, val))
    if seg_items and mode != "fotografico":
        add_section_title(doc, "Segurança", level=2)
        build_info_table(doc, seg_items)

    # ── Operacional (sem materiais - esses vão em tabela separada) ──
    op_items = []
    op_questions = [
        ("Capacidade operacional máxima da organização (ton/mês)", "Capacidade Operacional Máxima (t/mês)"),
        ("Quantidade de resíduos recebidos, comprados ou coletados da coleta convencional da prefeitura (TON)",
         "Resíduos da Coleta Convencional (t)"),
        ("Realiza a coleta e transporte de resíduos?", "Realiza Coleta e Transporte"),
        ("Realizam Pesagem?", "Realizam Pesagem"),
        ("Possuem ferramenta de controle de entrada/saída (planilha, software, etc)?",
         "Controle de Entrada/Saída"),
        ("Os destinatários possuem licença ambiental válida?", "Destinatários com Licença Ambiental"),
        ("Os resíduos não recicláveis são destinados a locais ambientalmente adequados após a triagem?",
         "Destinação Adequada de Não Recicláveis"),
        ("Comercializa créditos de logística reversa para outras entidades gestoras?",
         "Comercializa Créditos para Outras Entidades"),
    ]
    for q_key, label in op_questions:
        val = clean_answer(ans.get(q_key, ""))
        if val and val != "—":
            op_items.append((label, val))
    if op_items and mode != "fotografico":
        add_section_title(doc, "Operacional", level=2)
        build_info_table(doc, op_items)

    # ── Materiais Processados (tabela separada com mensal e anual) ──
    papel = ans.get("Papel", "")
    plastico = ans.get("Plástico", "")
    metal = ans.get("Metal", "")
    vidro = ans.get("Vidro", "")
    if any([papel, plastico, metal, vidro]) and mode != "fotografico":
        add_section_title(doc, "Materiais Processados", level=2)

        mat_values = [float(papel or 0), float(plastico or 0),
                      float(metal or 0), float(vidro or 0)]
        total_mes = sum(mat_values)
        mat_values_ano = [v * 12 for v in mat_values]
        total_ano = total_mes * 12

        mat_table = doc.add_table(rows=3, cols=6)
        mat_table.style = 'Table Grid'
        mat_table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        headers = ["", "Papel", "Plástico", "Metal", "Vidro", "TOTAL"]
        for i, h in enumerate(headers):
            cell = mat_table.cell(0, i)
            cell.text = ""
            set_cell_shading(cell, AZUL_HEX)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

        # Linha mensal
        row_mensal = ["t/mês"] + [f"{v:.0f}" for v in mat_values] + [f"{total_mes:.0f}"]
        for i, v in enumerate(row_mensal):
            cell = mat_table.cell(1, i)
            cell.text = ""
            if i == 0:
                set_cell_shading(cell, CINZA_CLARO)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_font(run, size=10, bold=(i == 0 or i == len(row_mensal) - 1))

        # Linha anual
        row_anual = ["t/ano"] + [f"{v:.0f}" for v in mat_values_ano] + [f"{total_ano:.0f}"]
        for i, v in enumerate(row_anual):
            cell = mat_table.cell(2, i)
            cell.text = ""
            if i == 0:
                set_cell_shading(cell, CINZA_CLARO)
            if i == len(row_anual) - 1:
                set_cell_shading(cell, CINZA_CLARO)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(v)
            set_font(run, size=10, bold=(i == 0 or i == len(row_anual) - 1))

        # Bordas
        tbl = mat_table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bn}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:color'), 'D0D0D0')
            borders.append(b)
        tblPr.append(borders)

        doc.add_paragraph()

    # ── Resíduos / Rejeitos ──
    waste_items = []
    rejeito_dest = ans.get("Qual o destino dos rejeitos?", "")
    rejeito_qty = ans.get("Quantidade Média de Rejeito (ton/mês)", "")
    if rejeito_dest:
        waste_items.append(("Destino dos Rejeitos", rejeito_dest))
    if rejeito_qty:
        waste_items.append(("Quantidade Média de Rejeito (t/mês)", rejeito_qty))
    if waste_items and mode != "fotografico":
        add_section_title(doc, "Resíduos / Rejeitos", level=2)
        build_info_table(doc, waste_items)

    # ── Contábil ──
    contabil_items = []
    for q_key, label in [("Possui Contador?", "Possui Contador"),
                          ("Utiliza ferramenta de controle financeiro?", "Controle Financeiro")]:
        val = clean_answer(ans.get(q_key, ""))
        if val and val != "—":
            contabil_items.append((label, val))
    if contabil_items and mode != "fotografico":
        add_section_title(doc, "Contábil", level=2)
        build_info_table(doc, contabil_items)

    # ── Observações (de todas as auditorias) ──
    if observations and mode != "fotografico":
        add_section_title(doc, "Observações", level=2)
        build_info_table(doc, observations)

    # ════════════════════════════════════════
    #  REGISTRO FOTOGRÁFICO
    # ════════════════════════════════════════
    if mode != "dco":
        doc.add_page_break()
        add_section_title(doc, "REGISTRO FOTOGRÁFICO", level=1)

        add_styled_paragraph(
            doc,
            "A seguir são apresentadas as fotos do empreendimento, "
            "organizadas por categoria de infraestrutura e operação.",
        )
        doc.add_paragraph()

        # Mapear fotos por category_name (excluir comprovante bancário e docs)
        fotos_by_cat = {}
        excluded_cats = {"Comprovante bancário", "Logotipo", "(sem categoria)"}
        for att in all_attachments:
            if att["category_name"] in excluded_cats:
                continue
            if att["category_group"] == "Documentation":
                continue
            fotos_by_cat.setdefault(att["category_name"], []).append(att)

        for cat_name, legend in PHOTO_SECTIONS:
            # Subtítulo da seção de foto
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(legend)
            set_font(run, size=13, bold=True, color=AZUL_ESCURO)

            photos = fotos_by_cat.pop(cat_name, [])
            if photos:
                for photo in photos:
                    add_photo_placeholder(doc, legend, photo=photo,
                                          fotos_dir=args.fotos_dir,
                                          storage_url=args.storage_url)
            else:
                add_photo_placeholder(doc, legend)

        # Extras
        for cat_name, photos in fotos_by_cat.items():
            if cat_name == "(sem categoria)":
                continue
            p = doc.add_paragraph()
            run = p.add_run(cat_name)
            set_font(run, size=13, bold=True, color=AZUL_ESCURO)
            for photo in photos:
                add_photo_placeholder(doc, cat_name, photo=photo,
                                      fotos_dir=args.fotos_dir,
                                      storage_url=args.storage_url)

    # ════════════════════════════════════════
    #  DECLARAÇÃO DE VÍNCULO TÉCNICO
    # ════════════════════════════════════════
    doc.add_page_break()
    add_section_title(doc, "DECLARAÇÃO DE VÍNCULO TÉCNICO", level=1)

    responsavel = args.responsavel or "Mariana Shiguemi Toschi Takagi"
    cpf_resp = args.cpf or "407.186.018-90"
    crea = args.crea or "CREA/SP 5070173382"

    add_styled_paragraph(
        doc,
        f"Eu, {responsavel}, CPF {cpf_resp}, declaro que fui a responsável técnica, "
        f"Engenheira Ambiental inscrita no {crea}, encarregada pela "
        f"POLEN CONSULTORIA E INTERMEDIAÇÃO DE NEGÓCIOS EM SUSTENTABILIDADE LTDA, "
        f"inscrita no CNPJ sobre o nº 28.038.406/0001-82, para realizar o preenchimento "
        f"da DECLARAÇÃO DE CAPACIDADE OPERACIONAL e RELATÓRIO FOTOGRÁFICO apresentados "
        f"acima do Operador de Reciclagem {company_name}, inscrita sob o CNPJ {cnpj}.",
    )

    for _ in range(3):
        doc.add_paragraph()

    hoje = datetime.now()
    data_fmt = f"{hoje.day} de {MESES_PT[hoje.month]} de {hoje.year}"
    add_styled_paragraph(doc, f"São Paulo, {data_fmt}.",
                         align=WD_ALIGN_PARAGRAPH.RIGHT)

    for _ in range(3):
        doc.add_paragraph()

    add_styled_paragraph(doc, "___________________________________________________",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, responsavel.upper(), size=11, bold=True,
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, "ENGENHEIRA AMBIENTAL",
                         align=WD_ALIGN_PARAGRAPH.CENTER)
    add_styled_paragraph(doc, crea, align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ═══════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="Gera relatório fotográfico Polen")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--subsidiary-id", help="UUID da subsidiary")
    group.add_argument("--cnpj", help="CNPJ do operador")
    group.add_argument("--name", help="Nome (parcial) do operador")
    parser.add_argument("--fotos-dir", help="Diretório com fotos baixadas")
    parser.add_argument("--storage-url", help="URL base do storage")
    parser.add_argument("--output", help="Arquivo de saída (.docx)")
    parser.add_argument("--responsavel", default="Mariana Shiguemi Toschi Takagi")
    parser.add_argument("--cpf", default="407.186.018-90")
    parser.add_argument("--crea", default="CREA/SP 5070173382")
    parser.add_argument("--mode", choices=["dco", "fotografico", "both"], default="both",
                        help="Tipo de documento: 'dco' (só DCO), 'fotografico' (só fotos) ou 'both' (legado)")
    args = parser.parse_args()

    print("Buscando operador...")
    operator = find_operator(args)
    print(f"  {operator['company_name']} | CNPJ: {operator['cnpj']} | {operator['type']}")

    print("Buscando auditoria mais recente...")
    audit = get_audit(operator["subsidiary_id"])
    aprovedat_str = str(audit.get('aprovedat', '—'))[:10]
    print(f"  ID: {audit['id'][:12]}... | Aprovada: {aprovedat_str}")

    print("Buscando fotos (todas as auditorias)...")
    all_attachments = get_all_attachments(operator["subsidiary_id"])
    print(f"  {len(all_attachments)} fotos")

    print("Buscando respostas do questionário...")
    answers = get_answers(audit["id"])
    print(f"  {len(answers)} respostas")

    print("Buscando observações...")
    observations = get_observations(operator["subsidiary_id"])
    print(f"  {len(observations)} observações")
    for label, val in observations:
        print(f"    - {label}: {val[:80]}")

    print("Gerando relatório...")
    doc = generate_report(operator, audit, all_attachments, answers, observations, args)

    if args.output:
        output_path = args.output
    else:
        name_clean = re.sub(r'[^\w]', '_', operator['company_name'][:40])
        output_path = f"Relatorio_Fotografico_{name_clean}.docx"

    doc.save(output_path)
    print(f"\nSalvo: {output_path}")


if __name__ == "__main__":
    main()
